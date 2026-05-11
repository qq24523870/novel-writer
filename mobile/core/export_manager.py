import os
import threading
from typing import Callable, List, Optional
from models.database import db_manager
from utils.logger import logger


class ExportManager:
    """导出管理器，支持多种格式导出"""

    @staticmethod
    def export_to_txt(chapters: List[dict], file_path: str,
                      on_progress: Callable[[int, int], None] = None) -> bool:
        """导出为TXT格式

        Args:
            chapters: 章节列表
            file_path: 导出路径
            on_progress: 进度回调

        Returns:
            是否成功
        """
        try:
            total = len(chapters)
            with open(file_path, "w", encoding="utf-8") as f:
                for i, chapter in enumerate(chapters):
                    f.write(f"\n{'='*50}\n")
                    f.write(f"{chapter['title']}\n")
                    f.write(f"{'='*50}\n\n")
                    f.write(chapter.get("content", ""))
                    f.write("\n\n")
                    if on_progress:
                        on_progress(i + 1, total)
            logger.info(f"导出TXT成功: {file_path}")
            return True
        except Exception as e:
            logger.error(f"导出TXT失败: {e}")
            return False

    @staticmethod
    def export_to_markdown(chapters: List[dict], file_path: str,
                           on_progress: Callable[[int, int], None] = None) -> bool:
        """导出为Markdown格式

        Args:
            chapters: 章节列表
            file_path: 导出路径
            on_progress: 进度回调

        Returns:
            是否成功
        """
        try:
            total = len(chapters)
            with open(file_path, "w", encoding="utf-8") as f:
                for i, chapter in enumerate(chapters):
                    f.write(f"\n# {chapter['title']}\n\n")
                    f.write(chapter.get("content", ""))
                    f.write("\n\n---\n\n")
                    if on_progress:
                        on_progress(i + 1, total)
            logger.info(f"导出Markdown成功: {file_path}")
            return True
        except Exception as e:
            logger.error(f"导出Markdown失败: {e}")
            return False

    @staticmethod
    def export_to_docx(chapters: List[dict], file_path: str,
                       on_progress: Callable[[int, int], None] = None) -> bool:
        """导出为Word (.docx) 格式

        Args:
            chapters: 章节列表
            file_path: 导出路径
            on_progress: 进度回调

        Returns:
            是否成功
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            style = doc.styles["Normal"]
            font = style.font
            font.name = "Microsoft YaHei"
            font.size = Pt(12)

            total = len(chapters)
            for i, chapter in enumerate(chapters):
                heading = doc.add_heading(chapter["title"], level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                content = chapter.get("content", "")
                paragraphs = content.split("\n")
                for para_text in paragraphs:
                    if para_text.strip():
                        p = doc.add_paragraph(para_text.strip())
                        p.paragraph_format.first_line_indent = Inches(0.28)
                        p.paragraph_format.line_spacing = 1.5

                doc.add_page_break()
                if on_progress:
                    on_progress(i + 1, total)

            doc.save(file_path)
            logger.info(f"导出DOCX成功: {file_path}")
            return True
        except ImportError:
            logger.error("导出DOCX失败: python-docx库未安装")
            return False
        except Exception as e:
            logger.error(f"导出DOCX失败: {e}")
            return False

    @staticmethod
    def export_async(export_format: str, chapters: List[dict], file_path: str,
                     on_complete: Callable[[bool], None] = None,
                     on_progress: Callable[[int, int], None] = None):
        """异步导出

        Args:
            export_format: 导出格式 (txt/markdown/docx)
            chapters: 章节列表
            file_path: 导出路径
            on_complete: 完成回调
            on_progress: 进度回调
        """
        def _run():
            if export_format == "txt":
                success = ExportManager.export_to_txt(chapters, file_path, on_progress)
            elif export_format == "markdown":
                success = ExportManager.export_to_markdown(chapters, file_path, on_progress)
            elif export_format == "docx":
                success = ExportManager.export_to_docx(chapters, file_path, on_progress)
            else:
                success = False

            if on_complete:
                on_complete(success)

        thread = threading.Thread(target=_run, daemon=True)
        thread.start()


export_manager = ExportManager()